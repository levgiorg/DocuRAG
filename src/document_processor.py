"""Document processing module for DocuRAG system.

Handles document ingestion, processing, and chunking for various file formats.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple

import chardet
try:
    import magic
except ImportError:
    magic = None
try:
    from loguru import logger
except ImportError:
    import logging
    logger = logging.getLogger(__name__)
from pydantic import BaseModel

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2 not available, PDF processing disabled")

try:
    import docx
except ImportError:
    docx = None
    logger.warning("python-docx not available, DOCX processing disabled")

from .utils import (
    calculate_file_hash,
    format_file_size,
    validate_file_extension,
    extract_metadata_from_filename,
    calculate_text_statistics,
    clean_text,
    timing_decorator
)


class DocumentChunk(BaseModel):
    """Represents a chunk of processed document."""
    
    content: str
    metadata: Dict[str, Any]
    chunk_index: int
    source_file: str
    page_number: Optional[int] = None
    start_char: Optional[int] = None
    end_char: Optional[int] = None
    quality_score: Optional[float] = None


class DocumentMetadata(BaseModel):
    """Document metadata container."""
    
    filename: str
    file_path: str
    file_hash: str
    file_size: int
    file_type: str
    processed_at: datetime
    title: Optional[str] = None
    author: Optional[str] = None
    creation_date: Optional[datetime] = None
    page_count: Optional[int] = None
    word_count: int = 0
    char_count: int = 0
    language: Optional[str] = None
    custom_metadata: Dict[str, Any] = {}


class DocumentProcessor:
    """Main document processing class."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize document processor.
        
        Args:
            chunk_size: Target size for text chunks
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = {'.pdf', '.txt', '.docx', '.md'}
        
        # Initialize file type detector
        if magic:
            try:
                self.magic = magic.Magic(mime=True)
            except Exception:
                self.magic = None
                logger.warning("python-magic not available, using extension-based detection")
        else:
            self.magic = None
            logger.info("python-magic not installed, using extension-based detection")
    
    def process_file(self, file_path: Union[str, Path]) -> Tuple[List[DocumentChunk], DocumentMetadata]:
        """Process a single file and return chunks with metadata.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Tuple[List[DocumentChunk], DocumentMetadata]: Processed chunks and metadata
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        logger.info(f"Processing file: {file_path.name}")
        
        # Extract text and metadata
        text_content, file_metadata = self._extract_content_and_metadata(file_path)
        
        # Clean and validate text
        text_content = clean_text(text_content)
        
        if not text_content.strip():
            logger.warning(f"No text content extracted from {file_path.name}")
            return [], file_metadata
        
        # Create chunks
        chunks = self._create_chunks(text_content, file_metadata)
        
        # Calculate quality scores
        self._calculate_quality_scores(chunks)
        
        logger.info(f"Created {len(chunks)} chunks from {file_path.name}")
        
        return chunks, file_metadata
    
    def process_directory(self, directory_path: Union[str, Path], 
                         recursive: bool = True) -> Tuple[List[DocumentChunk], List[DocumentMetadata]]:
        """Process all supported files in a directory.
        
        Args:
            directory_path: Directory to process
            recursive: Whether to process subdirectories
            
        Returns:
            Tuple[List[DocumentChunk], List[DocumentMetadata]]: All chunks and metadata
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        all_chunks = []
        all_metadata = []
        
        # Find files to process
        if recursive:
            files = [f for f in directory_path.rglob("*") if f.is_file()]
        else:
            files = [f for f in directory_path.iterdir() if f.is_file()]
        
        supported_files = [
            f for f in files 
            if f.suffix.lower() in self.supported_formats
        ]
        
        logger.info(f"Found {len(supported_files)} supported files in {directory_path}")
        
        for file_path in supported_files:
            try:
                chunks, metadata = self.process_file(file_path)
                all_chunks.extend(chunks)
                all_metadata.append(metadata)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        return all_chunks, all_metadata
    
    @timing_decorator
    def _extract_content_and_metadata(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Extract text content and metadata from file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple[str, DocumentMetadata]: Text content and metadata
        """
        file_extension = file_path.suffix.lower()
        
        # Base metadata
        base_metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_hash": calculate_file_hash(file_path),
            "file_size": file_path.stat().st_size,
            "processed_at": datetime.now()
        }
        
        # Add filename-based metadata
        base_metadata.update(extract_metadata_from_filename(file_path.name))
        
        # Extract content based on file type
        if file_extension == '.pdf':
            text_content, pdf_metadata = self._extract_pdf_content(file_path)
            base_metadata.update(pdf_metadata)
        elif file_extension == '.docx':
            text_content, docx_metadata = self._extract_docx_content(file_path)
            base_metadata.update(docx_metadata)
        elif file_extension in ['.txt', '.md']:
            text_content, txt_metadata = self._extract_text_content(file_path)
            base_metadata.update(txt_metadata)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Add text statistics
        text_stats = calculate_text_statistics(text_content)
        base_metadata.update(text_stats)
        base_metadata["file_type"] = file_extension
        
        metadata = DocumentMetadata(**base_metadata)
        
        return text_content, metadata
    
    def _extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple[str, Dict[str, Any]]: Text content and metadata
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        text_content = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    # Handle PDF date format
                    creation_date = pdf_reader.metadata.get('/CreationDate')
                    if creation_date and isinstance(creation_date, str):
                        # Convert PDF date format (D:20240410211143Z) to standard format
                        if creation_date.startswith('D:'):
                            date_str = creation_date[2:].replace('Z', '')
                            try:
                                from datetime import datetime
                                creation_date = datetime.strptime(date_str[:14], '%Y%m%d%H%M%S')
                            except ValueError:
                                creation_date = None
                    
                    metadata.update({
                        "title": pdf_reader.metadata.get('/Title'),
                        "author": pdf_reader.metadata.get('/Author'),
                        "creation_date": creation_date
                    })
                
                metadata["page_count"] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n\n--- Page {page_num} ---\n\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue
                
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            raise
        
        return text_content, metadata
    
    def _extract_docx_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple[str, Dict[str, Any]]: Text content and metadata
        """
        if docx is None:
            raise ImportError("python-docx is required for DOCX processing")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "creation_date": core_props.created
            }
            
            # Extract text content
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            raise
        
        return text_content, metadata
    
    def _extract_text_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Tuple[str, Dict[str, Any]]: Text content and metadata
        """
        metadata = {}
        
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'
                
            metadata["encoding"] = encoding
            metadata["encoding_confidence"] = encoding_result.get('confidence', 0.0)
            
            # Read text content
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                text_content = file.read()
                
        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {e}")
            raise
        
        return text_content, metadata
    
    def _create_chunks(self, text: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Create chunks from text content.
        
        Args:
            text: Text to chunk
            metadata: Document metadata
            
        Returns:
            List[DocumentChunk]: List of text chunks
        """
        # Split text into sentences for better chunking
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = ""
        current_length = 0
        chunk_index = 0
        start_char = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # Check if adding this sentence would exceed chunk size
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Create chunk
                chunk = self._create_chunk(
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    metadata=metadata,
                    start_char=start_char,
                    end_char=start_char + current_length
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + sentence
                current_length = len(overlap_text) + sentence_length
                start_char += current_length - len(overlap_text)
                chunk_index += 1
            else:
                current_chunk += sentence
                current_length += sentence_length
        
        # Add final chunk if there's content
        if current_chunk.strip():
            chunk = self._create_chunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                metadata=metadata,
                start_char=start_char,
                end_char=start_char + current_length
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for better chunking.
        
        Args:
            text: Input text
            
        Returns:
            List[str]: List of sentences
        """
        # Simple sentence splitting (can be improved with NLP libraries)
        sentence_pattern = r'(?<=[.!?])\s+'
        sentences = re.split(sentence_pattern, text)
        
        # Clean and filter sentences
        sentences = [s.strip() + " " for s in sentences if s.strip()]
        
        return sentences
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text for chunk continuity.
        
        Args:
            text: Current chunk text
            
        Returns:
            str: Overlap text
        """
        if len(text) <= self.chunk_overlap:
            return text
        
        # Take last chunk_overlap characters, but try to break at word boundary
        overlap_start = len(text) - self.chunk_overlap
        overlap_text = text[overlap_start:]
        
        # Find the first space to avoid breaking words
        space_index = overlap_text.find(' ')
        if space_index > 0:
            overlap_text = overlap_text[space_index:]
        
        return overlap_text
    
    def _create_chunk(self, content: str, chunk_index: int, metadata: DocumentMetadata,
                     start_char: int, end_char: int) -> DocumentChunk:
        """Create a DocumentChunk object.
        
        Args:
            content: Chunk content
            chunk_index: Index of the chunk
            metadata: Document metadata
            start_char: Start character position
            end_char: End character position
            
        Returns:
            DocumentChunk: Created chunk
        """
        chunk_metadata = {
            "source_document": metadata.filename,
            "document_hash": metadata.file_hash,
            "chunk_size": len(content),
            "created_at": datetime.now().isoformat()
        }
        
        # Add document metadata to chunk
        if metadata.title:
            chunk_metadata["document_title"] = metadata.title
        if metadata.author:
            chunk_metadata["document_author"] = metadata.author
        if metadata.page_count:
            chunk_metadata["total_pages"] = metadata.page_count
        
        return DocumentChunk(
            content=content,
            metadata=chunk_metadata,
            chunk_index=chunk_index,
            source_file=metadata.file_path,
            start_char=start_char,
            end_char=end_char
        )
    
    def _calculate_quality_scores(self, chunks: List[DocumentChunk]) -> None:
        """Calculate quality scores for chunks.
        
        Args:
            chunks: List of chunks to score
        """
        for chunk in chunks:
            score = self._calculate_chunk_quality(chunk.content)
            chunk.quality_score = score
    
    def _calculate_chunk_quality(self, text: str) -> float:
        """Calculate quality score for a text chunk.
        
        Args:
            text: Chunk text
            
        Returns:
            float: Quality score between 0 and 1
        """
        if not text.strip():
            return 0.0
        
        score = 0.5  # Base score
        
        # Length factor (prefer medium-length chunks)
        length_factor = min(len(text) / self.chunk_size, 1.0)
        score += 0.2 * length_factor
        
        # Word count factor
        words = text.split()
        if len(words) >= 10:
            score += 0.1
        
        # Sentence completeness (ends with punctuation)
        if text.strip()[-1] in '.!?':
            score += 0.1
        
        # Avoid chunks with mostly special characters or numbers
        alpha_ratio = sum(c.isalpha() for c in text) / len(text)
        score += 0.1 * alpha_ratio
        
        return min(score, 1.0)